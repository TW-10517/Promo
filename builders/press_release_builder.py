"""
Press Release builder: plain text -> .docx

The model returns the finished press release as plain text following the
template skeleton, so this builder only has to lay it out sensibly:

* the first non-empty line becomes the document heading
* lines starting with the section marker (■) become bold sub-headings
* everything else becomes a normal paragraph
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from core.missing_tracker import MissingItem

# Font that renders Japanese correctly on both Windows and macOS builds.
BODY_FONT = "Yu Gothic"
SECTION_MARKERS = ("■", "◆", "【")
_MISSING_MARKER_RE = re.compile(r"\[MISSING:\s*([^\]]*)\]")


def _set_base_style(document: Document) -> None:
    """Apply a readable base font to the whole document."""
    style = document.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    # python-docx needs the East-Asian font set explicitly for CJK runs.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        BODY_FONT,
    )


def detect_missing(text: str) -> list[MissingItem]:
    """
    Find every ``[MISSING: ...]`` marker in the raw press-release text.

    Tracks the most recent section heading (■/◆/【 line) the same way
    ``build()`` lays the document out, so ``context`` matches what a reader
    exporting the .docx would see. Shared by ``build()`` and the in-memory
    preview pipeline in ``core.router``.
    """
    missing: list[MissingItem] = []
    current_section: str | None = None
    for raw in text.replace("\r\n", "\n").split("\n"):
        line = raw.strip()
        if not line:
            continue
        if line.startswith(SECTION_MARKERS):
            current_section = line
        for match in _MISSING_MARKER_RE.finditer(line):
            field_name = match.group(1).strip() or "未指定項目"
            missing.append(MissingItem(field=field_name, context=current_section))
    return missing


def build(text: str, output_path: str | Path) -> tuple[Path, list[MissingItem]]:
    """
    Write the press release text to a .docx file.

    Parameters
    ----------
    text:
        Plain-text press release produced by Claude.
    output_path:
        Full destination path (including the .docx extension).

    Returns
    -------
    tuple[Path, list[MissingItem]]
        The written file, and every ``[MISSING: ...]`` marker found in the
        text -- content the Promotion Department still needs to confirm with
        the Planning Department. Generation is never blocked by this list;
        it is purely informational. ``context`` on each item is the section
        heading (■...) the marker appeared under, when there was one.
    """
    path = Path(output_path)
    missing = detect_missing(text)
    document = Document()
    _set_base_style(document)

    lines = text.replace("\r\n", "\n").split("\n")
    heading_done = False
    current_section: str | None = None

    for raw in lines:
        line = raw.rstrip()

        if not line.strip():
            # Preserve blank lines as spacing, but never start with one.
            if document.paragraphs:
                document.add_paragraph()
            continue

        # First meaningful line -> document title.
        if not heading_done:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line.strip())
            run.bold = True
            run.font.size = Pt(14)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading_done = True
            continue

        # Section markers -> bold sub-heading.
        if line.lstrip().startswith(SECTION_MARKERS):
            current_section = line.strip()
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line.strip())
            run.bold = True
            run.font.size = Pt(12)
            continue

        document.add_paragraph(line)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return path, missing
