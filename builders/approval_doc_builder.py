"""
Approval Document builder: JSON -> .docx

The JSON keys come straight from ``templates/approval_document.json``, so the
mapping is direct: header fields become a two-column table, sections become a
bold heading plus a body paragraph, and the approval stamp table is appended
at the end. Missing information is rendered as ``[MISSING: ...]`` in red so it
is impossible to miss during review.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from core.missing_tracker import MissingItem

BODY_FONT = "Yu Gothic"
MISSING_RE = re.compile(r"(\[MISSING:[^\]]*\])")
MISSING_COLOR = RGBColor(0xC0, 0x00, 0x00)  # dark red


def _set_base_style(document: Document) -> None:
    """Apply a readable base font (including the East-Asian font) document-wide."""
    style = document.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        BODY_FONT,
    )


def _add_value_runs(paragraph, value: str) -> None:
    """
    Write ``value`` into ``paragraph``, highlighting [MISSING: ...] markers.

    Literal ``\\n`` sequences and real newlines both become line breaks so a
    multi-line answer keeps its shape inside a single paragraph.
    """
    text = value.replace("\\n", "\n")
    first_line = True
    for line in text.split("\n"):
        if not first_line:
            paragraph.add_run().add_break()
        first_line = False
        for part in MISSING_RE.split(line):
            if not part:
                continue
            run = paragraph.add_run(part)
            if MISSING_RE.fullmatch(part):
                run.bold = True
                run.font.color.rgb = MISSING_COLOR


def detect_missing(data: dict, template: dict) -> list[MissingItem]:
    """
    Find every field that is blank or still carries a ``[MISSING: ...]`` marker.

    Runs against the same key/label pairs used to build the document, so
    detection can never drift from what was actually written to the file.
    Shared by ``build()`` (disk export) and the in-memory preview pipeline in
    ``core.router``, so both paths report the exact same missing fields.
    """
    items: list[MissingItem] = []
    for field in template.get("header_fields", []) + template.get("sections", []):
        value = str(data.get(field["key"], "")).strip()
        if not value or MISSING_RE.search(value):
            items.append(MissingItem(field=field["label"]))
    return items


def build(data: dict, template: dict, output_path: str | Path) -> tuple[Path, list[MissingItem]]:
    """
    Write the Approval Document.

    Parameters
    ----------
    data:
        Parsed JSON response, keyed by the template's field keys.
    template:
        The loaded ``approval_document.json`` structure.
    output_path:
        Full destination path (including the .docx extension).

    Returns
    -------
    tuple[Path, list[MissingItem]]
        The written file, and every field that came back blank or marked
        ``[MISSING: ...]`` -- content the Promotion Department still needs to
        confirm with the Planning Department. Generation is never blocked by
        this list; it is purely informational.
    """
    path = Path(output_path)
    missing = detect_missing(data, template)
    document = Document()
    _set_base_style(document)

    # --- Title -------------------------------------------------------------
    title = document.add_paragraph()
    title_run = title.add_run(template.get("document_title", "稟議書"))
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    # --- Header fields (two-column table) ----------------------------------
    header_fields = template.get("header_fields", [])
    if header_fields:
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for field in header_fields:
            row = table.add_row().cells
            label_paragraph = row[0].paragraphs[0]
            label_run = label_paragraph.add_run(field["label"])
            label_run.bold = True
            _add_value_runs(row[1].paragraphs[0], str(data.get(field["key"], "")).strip())
        # Keep the label column narrow.
        for row in table.rows:
            row.cells[0].width = Pt(110)
        document.add_paragraph()

    # --- Body sections -----------------------------------------------------
    for section in template.get("sections", []):
        heading = document.add_paragraph()
        heading_run = heading.add_run(section["label"])
        heading_run.bold = True
        heading_run.font.size = Pt(12)

        body = document.add_paragraph()
        _add_value_runs(body, str(data.get(section["key"], "")).strip())

    # --- Approval stamp table ---------------------------------------------
    approval = template.get("approval_table")
    if approval:
        document.add_paragraph()
        label = document.add_paragraph()
        label_run = label.add_run(approval.get("label", "承認欄"))
        label_run.bold = True
        label_run.font.size = Pt(12)

        columns = approval.get("columns", [])
        if columns:
            table = document.add_table(rows=2, cols=len(columns))
            table.style = "Table Grid"
            for index, column in enumerate(columns):
                cell = table.rows[0].cells[index]
                run = cell.paragraphs[0].add_run(column)
                run.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Second row is intentionally blank -- it is stamped/signed by hand.
            table.rows[1].height = Pt(48)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return path, missing
