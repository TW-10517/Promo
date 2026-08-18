"""
Generation router.

Owns the end-to-end pipeline for one run:
    PPT text (+ templates) -> prompt -> AI Client -> parsed response -> preview / file

Supports:
1. Preview generation in-memory (no local Word/Excel files created on disk)
2. Optional explicit file export when requested by the user
"""

from __future__ import annotations

import json
import shutil
from collections.abc import Callable
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from builders import action_plan_builder, approval_doc_builder, press_release_builder
from core import config, missing_tracker, prompts
from core.ai_client import AIClient
from core.missing_tracker import MissingItem

# Output identifiers
APPROVAL = "approval"
ACTION_PLAN = "action_plan"
PRESS_RELEASE = "press_release"

DISPLAY_NAMES_JA = {
    APPROVAL: "稟議書",
    ACTION_PLAN: "実施計画書",
    PRESS_RELEASE: "プレスリリース",
}

DISPLAY_NAMES_EN = {
    APPROVAL: "Approval Document",
    ACTION_PLAN: "Action Plan",
    PRESS_RELEASE: "Press Release",
}

DISPLAY_NAMES = DISPLAY_NAMES_JA

FILE_SUFFIXES = {
    APPROVAL: ("approval_document", ".docx"),
    ACTION_PLAN: ("action_plan", ".xlsx"),
    PRESS_RELEASE: ("press_release", ".docx"),
}

GENERATION_ORDER = [APPROVAL, ACTION_PLAN, PRESS_RELEASE]

ProgressFn = Callable[[str], None]


def _safe_stem(source_path: Path) -> str:
    stem = source_path.stem.strip() or "planning"
    for bad in '<>:"/\\|?*':
        stem = stem.replace(bad, "_")
    return stem[:60]


def _output_path(output_dir: Path, source_path: Path, kind: str, stamp: str) -> Path:
    name, ext = FILE_SUFFIXES[kind]
    base = f"{_safe_stem(source_path)}_{name}_{stamp}"
    candidate = output_dir / f"{base}{ext}"
    counter = 2
    while candidate.exists():
        candidate = output_dir / f"{base}_{counter}{ext}"
        counter += 1
    return candidate


@dataclass
class GeneratedDocument:
    """Represents an in-memory generated document with structured data and preview."""
    kind: str
    display_name_ja: str
    display_name_en: str
    raw_data: dict | str
    template: dict | str
    preview_text: str
    missing_items: list[MissingItem] = field(default_factory=list)
    # Filesystem-safe stem derived from the source .pptx, used as the default
    # filename if the user later chooses to export this document to disk.
    source_stem: str = "planning"
    # Set only for documents reconstructed from a real, already-exported
    # file (see build_legacy_document below). When set, Export copies this
    # file back out verbatim instead of trying to rebuild it from raw_data,
    # since structured raw_data can't be safely regenerated from plain text.
    source_file: Path | None = None


@dataclass
class RunResult:
    """Result of a generation run."""
    documents: dict[str, GeneratedDocument] = field(default_factory=dict)
    missing_by_output: dict[str, list[MissingItem]] = field(default_factory=dict)
    written: list[Path] = field(default_factory=list)
    files: dict[str, Path] = field(default_factory=dict)
    run_id: int | None = None
    selections: list[str] = field(default_factory=list)

    @property
    def total_missing(self) -> int:
        return sum(len(items) for items in self.missing_by_output.values())

    @property
    def missing_items(self) -> list[MissingItem]:
        return [item for items in self.missing_by_output.values() for item in items]

    @property
    def previews(self) -> dict[str, str]:
        return {doc.display_name_ja: doc.preview_text for doc in self.documents.values()}


# -- In-Memory Formatting Helpers ------------------------------------------

def _render_approval_preview(data: dict, template: dict) -> str:
    lines: list[str] = []
    lines.append("【稟議書 (社内決裁申請)】")
    lines.append("=" * 60)
    lines.append("")

    # Header
    for f_def in template.get("header_fields", []):
        val = str(data.get(f_def["key"], "[MISSING]")).strip()
        lines.append(f"● {f_def['label']} : {val}")
    lines.append("")
    lines.append("-" * 60)

    # Sections
    for sec in template.get("sections", []):
        val = str(data.get(sec["key"], "[MISSING]")).strip()
        lines.append("")
        lines.append(f"■ {sec['label']}")
        lines.append(val)
    lines.append("")

    # Approvers
    appr = template.get("approval_table", {})
    if appr:
        cols = appr.get("columns", [])
        lines.append("-" * 60)
        lines.append(f"【{appr.get('label', '承認欄')}】")
        lines.append("  |  " + "  |  ".join(cols) + "  |")
        lines.append("  |  " + "  |  ".join(["　　　" for _ in cols]) + "  |")

    return "\n".join(lines)


def _render_action_plan_preview(data: dict, template: dict) -> str:
    lines: list[str] = []
    lines.append("【実施計画書 (進行管理表)】")
    lines.append("=" * 60)
    lines.append("")

    rows_def = template.get("rows", [])
    cols_def = template.get("columns", [])

    if rows_def and cols_def:
        col_keys = [c.get("key", "") for c in cols_def]
        col_labels = [c.get("label", "") for c in cols_def]

        # Table header
        header_str = " | ".join(col_labels)
        lines.append(f"| {header_str} |")
        lines.append("|-" + "-|-".join(["-" * max(4, len(l)) for l in col_labels]) + "-|")

        # Rows
        rows_data = data.get("rows", [])
        for i, row_def in enumerate(rows_def):
            row_data = rows_data[i] if i < len(rows_data) else {}
            cells = []
            for ck in col_keys:
                val = str(row_data.get(ck, "")).strip()
                cells.append(val if val else "—")
            lines.append(f"| {' | '.join(cells)} |")
    else:
        lines.append(json.dumps(data, ensure_ascii=False, indent=2))

    return "\n".join(lines)


class Router:
    """Handles AI document generation pipelines."""

    def __init__(self, client: AIClient, progress: ProgressFn | None = None) -> None:
        self._client = client
        self._progress = progress or (lambda _message: None)

    def _log(self, message: str) -> None:
        self._progress(message)

    def _generate_approval(self, source_text: str) -> tuple[dict, dict]:
        template = config.load_json_template(config.TEMPLATE_APPROVAL)
        system, user = prompts.approval_document_prompt(source_text, template)
        schema = prompts.approval_document_schema(template)
        data = self._client.complete_json(system, user, schema)
        return data, template

    def _generate_action_plan(self, source_text: str, approval_text: str) -> tuple[dict, dict]:
        template = config.load_json_template(config.TEMPLATE_ACTION_PLAN)
        system, user = prompts.action_plan_prompt(source_text, approval_text, template)
        schema = prompts.action_plan_schema(template)
        data = self._client.complete_json(system, user, schema)
        return data, template

    def _generate_press_release(self, source_text: str) -> str:
        if config.get_provider() in {"local", "ollama"}:
            system, user, schema = prompts.press_release_json_prompt(source_text)
            data = self._client.complete_json(system, user, schema)
            return prompts.render_dospara_press_release(data)

        template = config.load_text_template(config.TEMPLATE_PRESS_RELEASE)
        system, user = prompts.press_release_prompt(source_text, template)
        return self._client.complete_text(system, user)

    def run_in_memory(
        self,
        source_text: str,
        selections: list[str],
        source_path: str | Path | None = None,
        db_path: Path | None = None,
    ) -> RunResult:
        """
        Generate documents purely in memory. No files are saved to disk.
        """
        stem = _safe_stem(Path(source_path)) if source_path else "planning"
        wanted = [k for k in GENERATION_ORDER if k in selections]
        documents: dict[str, GeneratedDocument] = {}
        collected_missing: list[MissingItem] = []

        approval_data: dict | None = None
        approval_template: dict | None = None

        for kind in wanted:
            if kind == APPROVAL:
                self._log("稟議書を生成中...")
                approval_data, approval_template = self._generate_approval(source_text)
                preview = _render_approval_preview(approval_data, approval_template)

                missing = approval_doc_builder.detect_missing(approval_data, approval_template)
                for item in missing:
                    item.output = DISPLAY_NAMES_JA[APPROVAL]
                collected_missing.extend(missing)

                documents[APPROVAL] = GeneratedDocument(
                    kind=APPROVAL,
                    display_name_ja=DISPLAY_NAMES_JA[APPROVAL],
                    display_name_en=DISPLAY_NAMES_EN[APPROVAL],
                    raw_data=approval_data,
                    template=approval_template,
                    preview_text=preview,
                    missing_items=missing,
                    source_stem=stem,
                )
                self._log("✓ 稟議書の生成が完了しました")

            elif kind == ACTION_PLAN:
                if approval_data is None:
                    self._log("実施計画書の生成に必要な稟議書データを準備中...")
                    approval_data, approval_template = self._generate_approval(source_text)

                approval_text = prompts.flatten_approval(approval_data, approval_template or {})
                self._log("実施計画書を生成中...")
                data, template = self._generate_action_plan(source_text, approval_text)
                preview = _render_action_plan_preview(data, template)

                missing = action_plan_builder.detect_missing(data, template)
                for item in missing:
                    item.output = DISPLAY_NAMES_JA[ACTION_PLAN]
                collected_missing.extend(missing)

                documents[ACTION_PLAN] = GeneratedDocument(
                    kind=ACTION_PLAN,
                    display_name_ja=DISPLAY_NAMES_JA[ACTION_PLAN],
                    display_name_en=DISPLAY_NAMES_EN[ACTION_PLAN],
                    raw_data=data,
                    template=template,
                    preview_text=preview,
                    missing_items=missing,
                    source_stem=stem,
                )
                self._log("✓ 実施計画書の生成が完了しました")

            elif kind == PRESS_RELEASE:
                self._log("プレスリリースを生成中...")
                text = self._generate_press_release(source_text)

                missing = press_release_builder.detect_missing(text)
                for item in missing:
                    item.output = DISPLAY_NAMES_JA[PRESS_RELEASE]
                collected_missing.extend(missing)

                documents[PRESS_RELEASE] = GeneratedDocument(
                    kind=PRESS_RELEASE,
                    display_name_ja=DISPLAY_NAMES_JA[PRESS_RELEASE],
                    display_name_en=DISPLAY_NAMES_EN[PRESS_RELEASE],
                    raw_data=text,
                    template="",
                    preview_text=text,
                    missing_items=missing,
                    source_stem=stem,
                )
                self._log("✓ プレスリリースの生成が完了しました")

        missing_by_output: dict[str, list[MissingItem]] = {}
        if db_path and collected_missing:
            try:
                run_id = missing_tracker.create_run(
                    source_filename="memory",
                    timestamp=datetime.now(),
                    db_path=db_path,
                )
                missing_tracker.add_missing_items(run_id, collected_missing, db_path=db_path)
            except Exception:
                pass  # best-effort persistence; items are still shown, just without resolvable ids
        for item in collected_missing:
            missing_by_output.setdefault(item.output, []).append(item)

        return RunResult(
            documents=documents,
            missing_by_output=missing_by_output,
            selections=list(wanted),
        )

    @staticmethod
    def export_document(doc: GeneratedDocument, target_path: Path) -> Path:
        """
        Export a single generated document to Word (.docx) or Excel (.xlsx).

        Pure file-writing (no AI call), so it's a staticmethod: exporting a
        document that already exists in memory should never require a live
        AI client or a configured API key.
        """
        target_path.parent.mkdir(parents=True, exist_ok=True)
        if doc.source_file is not None and doc.source_file.is_file():
            shutil.copy2(doc.source_file, target_path)
            return target_path
        if doc.kind == APPROVAL:
            path, _ = approval_doc_builder.build(doc.raw_data, doc.template, target_path)
            return path
        elif doc.kind == ACTION_PLAN:
            path, _ = action_plan_builder.build(doc.raw_data, doc.template, target_path)
            return path
        elif doc.kind == PRESS_RELEASE:
            path, _ = press_release_builder.build(doc.raw_data, target_path)
            return path
        raise ValueError(f"Unknown document kind: {doc.kind}")


# ---------------------------------------------------------------------------
# Internal generation state persistence
#
# Saves the raw generated text/data (never a .docx/.xlsx) into the project's
# own folder so reopening a project can show its previous output again. This
# is separate from ``export_document`` above, which is the only code path
# that ever writes a real Word/Excel file.
# ---------------------------------------------------------------------------

def save_state(state_path: Path, result: RunResult, selections: list[str]) -> None:
    """Persist every generated document's raw text/data to disk as JSON."""
    state_path.parent.mkdir(parents=True, exist_ok=True)
    payload = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "selections": selections,
        "documents": {
            kind: {
                "kind": doc.kind,
                "display_name_ja": doc.display_name_ja,
                "display_name_en": doc.display_name_en,
                "raw_data": doc.raw_data,
                "template": doc.template,
                "preview_text": doc.preview_text,
                "source_stem": doc.source_stem,
                "missing_items": [
                    {
                        "field": item.field,
                        "context": item.context,
                        "output": item.output,
                        "id": item.id,
                        "resolved": item.resolved,
                        "resolved_at": item.resolved_at,
                    }
                    for item in doc.missing_items
                ],
            }
            for kind, doc in result.documents.items()
        },
    }
    state_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def load_state(state_path: Path) -> RunResult | None:
    """Load a previously generated run for this project, if one exists."""
    if not state_path.is_file():
        return None
    try:
        payload = json.loads(state_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None

    documents: dict[str, GeneratedDocument] = {}
    missing_by_output: dict[str, list[MissingItem]] = {}
    for kind, d in payload.get("documents", {}).items():
        items = [
            MissingItem(
                field=it["field"],
                context=it.get("context"),
                output=it.get("output", ""),
                id=it.get("id"),
                resolved=it.get("resolved", False),
                resolved_at=it.get("resolved_at"),
            )
            for it in d.get("missing_items", [])
        ]
        documents[kind] = GeneratedDocument(
            kind=d["kind"],
            display_name_ja=d["display_name_ja"],
            display_name_en=d["display_name_en"],
            raw_data=d["raw_data"],
            template=d["template"],
            preview_text=d["preview_text"],
            missing_items=items,
            source_stem=d.get("source_stem", "planning"),
        )
        for item in items:
            missing_by_output.setdefault(item.output, []).append(item)

    return RunResult(
        documents=documents,
        missing_by_output=missing_by_output,
        selections=payload.get("selections", []),
    )


# ---------------------------------------------------------------------------
# Legacy document support
#
# Projects generated before internal-state tracking existed have real
# exported .docx/.xlsx files but no JSON state to reload. Rather than ever
# handing off to an external app (Word/Excel) to view them, read their text
# back in and wrap it as a normal GeneratedDocument, so they open in the
# same in-app notepad view as anything generated today.
# ---------------------------------------------------------------------------

def read_legacy_document_text(path: Path) -> str:
    """Read a previously exported .docx/.xlsx back into plain text."""
    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            from docx import Document

            document = Document(str(path))
            lines = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        lines.append(" | ".join(cells))
            return "\n".join(lines).strip()

        if suffix == ".xlsx":
            from openpyxl import load_workbook

            workbook = load_workbook(str(path), data_only=True)
            lines = []
            for sheet in workbook.worksheets:
                lines.append(f"[{sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    cells = ["" if c is None else str(c) for c in row]
                    if any(cells):
                        lines.append(" | ".join(cells))
            return "\n".join(lines).strip()

        return path.read_text(encoding="utf-8")
    except Exception as exc:
        return f"[このファイルの内容を読み込めませんでした: {exc}]"


def build_legacy_document(path: Path) -> GeneratedDocument:
    """Wrap a real, previously exported file as an in-app-viewable document."""
    name = path.name
    if "action_plan" in name.lower() or "実施計画" in name:
        kind = ACTION_PLAN
    elif "press_release" in name.lower() or "プレスリリース" in name:
        kind = PRESS_RELEASE
    else:
        kind = APPROVAL

    text = read_legacy_document_text(path)
    return GeneratedDocument(
        kind=kind,
        display_name_ja=DISPLAY_NAMES_JA[kind],
        display_name_en=DISPLAY_NAMES_EN[kind],
        raw_data=text,
        template="",
        preview_text=text,
        missing_items=[],
        source_stem=path.stem,
        # Structured kinds can't be safely rebuilt from plain extracted
        # text, so Export just copies this real file back out verbatim.
        # Press releases are already plain text, so they export normally
        # (edits included) without needing this.
        source_file=path if kind in (APPROVAL, ACTION_PLAN) else None,
    )
