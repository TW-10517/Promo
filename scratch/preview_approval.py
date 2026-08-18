import sys
from pathlib import Path
from docx import Document

doc_path = Path("projects/GALLERIA_Autumn_2026_Promotion_20260814_172306_cb7bf3/output/demo_planning_document_approval_document_20260814-172512.docx")
if not doc_path.is_file():
    print("File not found")
    sys.exit(0)

doc = Document(str(doc_path))
out = []
out.append("=== GENERATED APPROVAL DOCUMENT PREVIEW ===")
for p in doc.paragraphs:
    if p.text.strip():
        out.append(p.text.strip())

for t in doc.tables:
    out.append("\n--- HEADER TABLE ---")
    for r in t.rows:
        row_str = " | ".join(c.text.strip().replace("\n", " ") for c in r.cells)
        out.append(row_str)

sys.stdout.buffer.write(("\n".join(out) + "\n").encode("utf-8"))
